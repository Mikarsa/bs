from docx import Document
import re
import xml.etree.ElementTree as ET
import os
import win32com.client as win32

def convert_doc_to_docx(doc_path):
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(doc_path + "x", FileFormat=16)  # 16 表示 .docx 格式
    doc.Close()
    word.Quit()


def remove_patterns(text):
    pattern = r'GB_\.+|S__'  # 匹配 GB_后任意两字符 或 S__
    return re.sub(pattern, '', text).strip()


def docx_input_gb(file_path):
    l = {}
    str1 = ''
    str2 = ''
    doc = Document(file_path)
    flag = False
    for paragraph in doc.paragraphs:
        data = paragraph.text.replace(' ', '')
        # print(data)
        # print(paragraph.text)
        if not flag:
            # 匹配标题
            pattern = r'\d+\.\d+\.\d+\.\d+(.*)'
            match = re.search(pattern, data)
            if match:
                flag = True
                str1 = match.group(0)
        elif data.find('漏洞描述：') != -1:
            str2 = data[data.find('漏洞描述：') + len('漏洞描述：') :]
            l[str2]=str1
            flag = False
            print(str2+' '+str1)
    return l


def docx_input_zbg(file_path):
    l = []
    doc = Document(file_path)
    for table in doc.tables:
        flag = 0
        # print(table)
        for row in table.rows:  # 遍历每一行
            for cell in row.cells:  # 便利每个单元格
                if cell.text == '漏洞类别':
                    flag = 1
                if flag == 0:
                    continue
                txt = (cell.text.translate(str.maketrans('', '', '0123456789')))
                l.append(txt)
                # print(cell.text)
    return list(i for i in set(l) if i.strip())


def xml_imput_kubo(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    problem_summary = root.find('你的问题汇总')
    rules_totals = problem_summary.findall('ruleTotal')
    result = []
    for rule_total in rules_totals:
        rule_total_name = rule_total.get('规则集')

        for rule in rule_total.findall('rule'):
            rule_data = {
                "序号": rule.get("序号"),
                "名称": rule.get("名称"),
                "严重等级": rule.get("严重等级"),
                "数量": rule.get("数量")
            }
            # result.append(''.join(re.findall('[\u4e00-\u9fa5]',rule_data["名称"])))
            result.append(rule_data["名称"].split(' ', 1)[1])
            # print(rule_data["名称"].split(' ', 1)[1])
            # print(rule_data)
    # print(len(result))
    # print(result)
    return result


def docx_input_keda(file_path):
    l = []
    doc = Document(file_path)
    for table in doc.tables:
        flag = 0
        # print(table)
        for row in table.rows:  # 遍历每一行
            for cell in row.cells:  # 便利每个单元格
                if cell.text == '序号':
                    flag = 1
                if flag == 0:
                    continue
                txt = (cell.text.translate(str.maketrans('', '', '0123456789')))

                if txt not in ['重大', '严重', '次要', '序号', '重要', '数量', '严重等级', '名称', '合计']:
                    l.append(remove_patterns(txt))
                # print(cell.text)
    return list(i for i in set(l) if i.strip())


def docx_append_period(file_path):
    doc = Document(file_path)
    for para in doc.paragraphs:
        if para.text.find('。') != -1:
            pass



def docx_input_gbsy(file_path):
    a = 0
    l = []
    doc = Document(file_path)
    tag = False
    for para in doc.paragraphs:
        if tag:
            temp = para.text[:para.text.find(' ')]
            if len(temp) != 1: # 去除长度只有一的词以及标点符号
                l.append(temp)
            tag = False
            continue
        pattern = r'^\d+\.\d+$'
        txt = para.text.replace(' ', '')
        if bool(re.fullmatch(pattern, txt)) or bool(re.fullmatch(pattern, txt[txt.rfind('。')+1:])):
            if(txt.find('。') == -1):
                temp = txt
                # l.append(txt)
            else:
                temp = txt[txt.rfind('。')+1:]
                # l.append(txt[txt.rfind('。')+1:])
            tag = True
            a += 1
            if a != int(temp[temp.find('.')+1:]):
                print(a)
                a+=1
            # l.append(temp)
    return l


def output_txt(file_path, data):
    with open(file_path, 'w', encoding='utf-8') as f:
        for i in data:
            f.write(i + '\n')

if __name__ == '__main__':
    origin_path = 'E:/bs/Database/'
    # gb提取数据集
    #dx = docx_input_gb('E:/codespace/GBT34943(1).docx')

    # zbg提取数据集
    # l = []
    # for i in range(1,501):
    #     dx = docx_input_zbg(f'E:/database/lrj-report/1/lrj{i}-综合审计报告.docx')
    #     l += dx
    #     # print(f'{i}:\n{dx}')
    # for i in range(501, 600):
    #     dx = docx_input_zbg(f'E:/database/lrj-report/2/lrj{i}-综合审计报告.docx')
    #     print(f'{i}:\n{dx}')
    #     l += dx
    #
    # with open('E:/codespace/report_text.txt', 'w', encoding='utf-8') as f:
    #     for i in list(set(l)):
    #         if i in ['危险', '高', '中', '低', '漏洞数', '漏洞类别']:
    #             continue
    #         f.write(i + '\n')
    # print(list(set(l)))

    # 库博报告提取数据
    # report = []
    # xml_kubo_filepath = origin_path + '库博项目报告'
    # xml_kubo_file = [ xml_kubo_filepath + '/' + i for i in os.listdir(xml_kubo_filepath) if i.find('.xml') != -1]
    # a = 0
    # for i in xml_kubo_file:
    #     report += xml_imput_kubo(i)
    #
    #
    # report = list(set(report))
    # output_txt(origin_path + 'kubo.txt', report)
    #
    # # 科大报告提取数据
    # # convert_doc_to_docx(origin_path + '科大项目报告/1000-2-报告/1000-2-检测报告-概述汇总.doc')
    # file_path_keda = origin_path  + '科大项目报告/1000-报告/1000-检测报告-概述汇总.docx'
    # report = docx_input_keda(file_path_keda)
    # file_path_keda1 = origin_path  + '科大项目报告/1000-2-报告/1000-2-检测报告-概述汇总.docx'
    # report += docx_input_keda(file_path_keda1)
    # report = list(set(report))
    # # print(len(report), report)
    #
    # output_txt(origin_path + 'keda.txt', report)

    # gbsy报告提取数据
    file_path_gbsy = origin_path + 'GBT+25069-2022.docx'
    text = docx_input_gbsy(file_path_gbsy)
    output_txt(origin_path + 'GBT+25069-2022.txt', text)

